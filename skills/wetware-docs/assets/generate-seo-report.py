#!/usr/bin/env python3
"""
Generate an empty Wetware Labs SEO Audit Report template.
Uses the official Wetware_Labs_Template.docx for header/footer.

Output: A .docx with all sections, placeholder text, and styling ready
for any client. Fill in the data and generate.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from copy import deepcopy
import io
import os
from zipfile import ZipFile
from datetime import datetime

SKILL_DIR = "/Users/myrm/obsidian/vaults/.claude/skills/wetware-docs"
TEMPLATE_PATH = os.path.join(SKILL_DIR, "assets", "Wetware_Labs_Template.docx")
LOGO_PATH = os.path.join(SKILL_DIR, "assets", "wetwareArtboard 1@4x.png")
OUTPUT_DIR = "/Users/myrm/Desktop"
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "Wetware_SEO_Report_Template.docx")

# ── Colors ────────────────────────────────────────────────────────────────────
BLACK = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ── Helpers ───────────────────────────────────────────────────────────────────

def add_text(doc, text, size=10, color=DARK_GRAY, bold=False, italic=False,
             align=None, space_before=0, space_after=0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p


def section_heading(doc, text, page_break=False):
    p = doc.add_paragraph()
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.color.rgb = BLACK
    run.bold = True
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '3',
        qn('w:space'): '3', qn('w:color'): '000000'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.color.rgb = BLACK
    run.bold = True
    return p


def add_spacer(doc, space_pt=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_pt)
    p.paragraph_format.space_after = Pt(0)
    return p


def set_cell_shading(cell, color_hex):
    shading = cell._tc.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex
    })
    shading.append(shd)


def set_cell_text(cell, text, size=10, color=DARK_GRAY, bold=False,
                  align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = tcBorders.makeelement(qn(f'w:{side}'), {
                qn('w:val'): 'single', qn('w:sz'): val.get('sz', '4'),
                qn('w:space'): '0', qn('w:color'): val.get('color', '000000')
            })
            tcBorders.append(border)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = borders.makeelement(qn(f'w:{side}'), {
            qn('w:val'): 'none', qn('w:sz'): '0',
            qn('w:space'): '0', qn('w:color'): 'auto'
        })
        borders.append(border)
    tblPr.append(borders)


def issue_row(doc, issue_text, description, count, pct=None):
    """Create a single issue row: issue name, description, count, percentage."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(5.2)
    table.columns[1].width = Inches(1.3)
    remove_table_borders(table)

    # Bottom border on cells
    for col in range(2):
        set_cell_borders(table.cell(0, col), bottom={'sz': '2', 'color': 'CCCCCC'})

    # Left cell: issue name + description
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(issue_text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    run.bold = True

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(description)
    run2.font.name = "Arial"
    run2.font.size = Pt(8)
    run2.font.color.rgb = GRAY
    run2.italic = True

    # Right cell: count
    right_cell = table.cell(0, 1)
    right_cell.text = ""
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(str(count))
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = BLACK
    run.bold = True


# ── Build the document ────────────────────────────────────────────────────────

doc = Document(TEMPLATE_PATH)

# Clear body, preserve sectPr
body = doc.element.body
sect_pr = body.findall(qn('w:sectPr'))
sect_pr_copy = [deepcopy(sp) for sp in sect_pr]
for child in list(body):
    body.remove(child)
for sp in sect_pr_copy:
    body.append(sp)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1: COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

# Big spacer to push content down
for _ in range(6):
    add_spacer(doc, 24)

# WETWARE logo (centered, large)
logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_p.paragraph_format.space_after = Pt(0)
logo_run = logo_p.add_run()
logo_run.add_picture(LOGO_PATH, width=Inches(4.0))

# More space
for _ in range(3):
    add_spacer(doc, 24)

# Report title
add_text(doc, "Full Site Audit Report", size=26, color=BLACK, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

# Domain placeholder
add_text(doc, "[client-domain.com]", size=12, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

# Push date to bottom
for _ in range(6):
    add_spacer(doc, 24)

# Generated date
add_text(doc, f"Generated on [Date]", size=10, color=LIGHT_GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 2: EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

section_heading(doc, "Executive Summary", page_break=True)

add_spacer(doc, 6)

# Score metrics table (borderless, single row)
metrics_table = doc.add_table(rows=2, cols=5)
metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(metrics_table)

metric_labels = ["Site Health", "Pages Crawled", "Errors", "Warnings", "Notices"]
metric_placeholders = ["___%", "___", "___", "___", "___"]

for i in range(5):
    # Big number
    set_cell_text(metrics_table.cell(0, i), metric_placeholders[i],
                  size=28, color=BLACK, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    # Label
    set_cell_text(metrics_table.cell(1, i), metric_labels[i],
                  size=8, color=GRAY, bold=False,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

add_spacer(doc, 12)

# Page Breakdown
sub_heading(doc, "Page Breakdown")
add_text(doc, "[Insert page breakdown bar: Healthy (___) | Broken (___) | Have Issues (___) | Redirected (___) | Blocked (___)]",
         size=9, color=GRAY, italic=True, space_after=8)

# Key Findings
sub_heading(doc, "Key Findings")

findings_placeholders = [
    "[___% Site Health score indicates ___. Summary of overall health.]",
    "[___ broken internal links are the most critical issue (___ % of all problems). Brief description.]",
    "[___ pages return 4XX errors and ___ have duplicate title tags, both hurting search visibility.]",
    "[___ pages have low text-to-HTML ratio — search engines prefer content-heavy pages over code-heavy ones.]",
    "[___ orphaned pages exist in sitemaps but have no internal links pointing to them, making them invisible to visitors.]",
]

for finding in findings_placeholders:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"\u2022  {finding}")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GRAY

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 3: TOP ISSUES
# ═══════════════════════════════════════════════════════════════════════════════

section_heading(doc, "Top Issues", page_break=True)

top_issues = [
    ("[___ internal links are broken]", "[Progress bar placeholder]", "___", "___% of total issues"),
    ("[___ pages returned 4XX status code]", "[Progress bar placeholder]", "___", "___% of total issues"),
    ("[___ issues with duplicate title tags]", "[Progress bar placeholder]", "___", "___% of total issues"),
    ("[___ pages have duplicate meta descriptions]", "[Progress bar placeholder]", "___", "___% of total issues"),
    ("[___ pages have low text-to-HTML ratio]", "[Progress bar placeholder]", "___", "___% of total issues"),
    ("[___ pages have too much text in title tags]", "[Progress bar placeholder]", "___", "___% of total issues"),
    ("[___ pages don't have an h1 heading]", "[Progress bar placeholder]", "___", "___% of total issues"),
]

for issue_text, _bar, count, _pct in top_issues:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(5.2)
    table.columns[1].width = Inches(1.3)
    remove_table_borders(table)
    set_cell_borders(table.cell(0, 0), bottom={'sz': '2', 'color': 'CCCCCC'})
    set_cell_borders(table.cell(0, 1), bottom={'sz': '2', 'color': 'CCCCCC'})

    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"\u25A0  {issue_text}")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

    right_cell = table.cell(0, 1)
    right_cell.text = ""
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(_pct)
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 4: ERRORS
# ═══════════════════════════════════════════════════════════════════════════════

section_heading(doc, "Errors \u2014 ___ Total", page_break=True)

error_items = [
    ("Internal links are broken", "Fix or remove all broken internal links. Visitors clicking these reach dead pages.", "___"),
    ("Pages returned 4XX status code", "Remove links to these pages or restore the content. 404 errors hurt crawlability.", "___"),
    ("Duplicate title tags", "Write a unique <title> for each page with relevant keywords.", "___"),
    ("Duplicate meta descriptions", "Write unique meta descriptions for each page.", "___"),
]

for name, desc, count in error_items:
    issue_row(doc, name, desc, count)

# ═══════════════════════════════════════════════════════════════════════════════
# WARNINGS
# ═══════════════════════════════════════════════════════════════════════════════

section_heading(doc, "Warnings \u2014 ___ Total")

warning_items = [
    ("Low text-to-HTML ratio", "Add more meaningful text content. Aim for 25%+ text-to-code ratio.", "___"),
    ("Title tags too long", "Keep page titles under 70 characters so they display fully in search results.", "___"),
    ("Missing h1 heading", "Add a single, descriptive <h1> tag to every page.", "___"),
    ("Missing meta descriptions", "Add meta descriptions containing keywords to improve click-through rates.", "___"),
    ("Low word count", "Pages with fewer than 200 words signal thin content. Expand with useful information.", "___"),
    ("Broken external link", "Fix or remove the broken outbound link.", "___"),
]

for name, desc, count in warning_items:
    issue_row(doc, name, desc, count)

# ═══════════════════════════════════════════════════════════════════════════════
# NOTICES
# ═══════════════════════════════════════════════════════════════════════════════

section_heading(doc, "Notices \u2014 ___ Total", page_break=True)

notice_items = [
    ("Pages require content optimization", "Optimize content for clarity, keyword usage, and structured headings.", "___"),
    ("Orphaned pages in sitemaps", "Add internal links to orphaned pages or remove them from the sitemap.", "___"),
    ("Multiple h1 tags per page", "Use only one <h1> per page. Use <h2>\u2013<h6> for sub-sections.", "___"),
    ("Permanent redirects (301/308)", "Review all redirects and update links to point to final URLs.", "___"),
    ("Pages with only one internal link", "Improve internal linking to boost discoverability and authority.", "___"),
    ("Pages blocked from crawling", "Verify blocked pages are intentionally excluded from search.", "___"),
    ("Llms.txt not found", "Create an llms.txt file to help AI search engines understand your site.", "___"),
]

for name, desc, count in notice_items:
    issue_row(doc, name, desc, count)


# ═══════════════════════════════════════════════════════════════════════════════
# SQUARESPACE LIMITATIONS & PLATFORM CEILING
# ═══════════════════════════════════════════════════════════════════════════════

section_heading(doc, "What Can \u2014 and Can\u2019t \u2014 Be Fixed on Squarespace", page_break=True)

add_text(doc, "Squarespace is a managed website platform. It handles hosting, SSL, and basic "
         "infrastructure so you don\u2019t have to. But that convenience comes with hard ceilings "
         "on what\u2019s technically possible. Some of the issues in this report can be fully resolved "
         "within Squarespace. Others are permanent platform constraints that would only be solved "
         "by migrating to a custom-built site.",
         size=10, color=DARK_GRAY, space_before=6, space_after=12)

# Estimated ceiling
sub_heading(doc, "Estimated Maximum SEO Health on Squarespace: ___% (currently ___%)")

add_text(doc, "Even with every fixable issue resolved, Squarespace\u2019s built-in code bloat, "
         "limited heading control, and restricted file access prevent a site from reaching the "
         "same technical SEO score as a custom-built site. The items below explain exactly what "
         "that ceiling looks like and how it affects your ability to rank.",
         size=9, color=GRAY, italic=True, space_before=2, space_after=16)

# ── FIXABLE ON SQUARESPACE ──
sub_heading(doc, "\u2705  Fixable Within Squarespace")

fixable_items = [
    ("Title tags (duplicate & too long)",
     "Squarespace lets you set a unique SEO title per page under Page Settings \u2192 SEO. "
     "Write keyword-rich titles under 70 characters for every page. This resolves duplicate "
     "titles and truncation in search results.",
     "High"),
    ("Meta descriptions (duplicate & missing)",
     "Each page has a meta description field in Page Settings \u2192 SEO. Write unique, "
     "compelling descriptions under 160 characters with your target keywords.",
     "High"),
    ("Broken internal links",
     "Most broken links are phone numbers formatted as URLs instead of tel: links. "
     "Edit each phone number link to use the format tel:+1XXXXXXXXXX. Also fix or remove "
     "any links pointing to deleted or renamed pages.",
     "Critical"),
    ("4XX errors (broken pages)",
     "Either restore deleted pages, create URL redirects (Settings \u2192 URL Mappings), "
     "or remove internal links pointing to dead URLs.",
     "High"),
    ("Low word count / thin content",
     "Add meaningful, keyword-relevant text to pages with fewer than 200 words. "
     "Squarespace\u2019s editor makes this straightforward \u2014 add text blocks to any page.",
     "Medium"),
    ("Broken external links",
     "Find and update or remove any outbound links that point to dead external URLs.",
     "Low"),
    ("Orphaned pages in sitemaps",
     "Add internal links from your main pages to orphaned pages, or hide them from "
     "navigation if they\u2019re not needed. Squarespace auto-generates sitemaps, so the "
     "fix is adding links, not editing the sitemap directly.",
     "Medium"),
    ("Pages with only one internal link",
     "Add contextual internal links between related service pages. Link your blog posts "
     "to service pages. Link service pages to your contact page. Every important page "
     "should be reachable within 3 clicks from the homepage.",
     "Medium"),
    ("Content optimization",
     "Restructure page content with clear headings (H2, H3), target keywords in the "
     "first paragraph, and add FAQ sections where relevant.",
     "Medium"),
    ("Permanent redirects (301/308)",
     "Set up URL redirects in Settings \u2192 URL Mappings to point old URLs to current ones. "
     "This consolidates link equity and stops crawlers from hitting dead ends.",
     "Medium"),
]

for name, desc, priority in fixable_items:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(1.0)
    remove_table_borders(table)
    set_cell_borders(table.cell(0, 0), bottom={'sz': '2', 'color': 'DDDDDD'})
    set_cell_borders(table.cell(0, 1), bottom={'sz': '2', 'color': 'DDDDDD'})

    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"\u2705  {name}")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    run.bold = True

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(desc)
    run2.font.name = "Arial"
    run2.font.size = Pt(8.5)
    run2.font.color.rgb = DARK_GRAY

    right_cell = table.cell(0, 1)
    right_cell.text = ""
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    priority_color = BLACK if priority == "Critical" else DARK_GRAY if priority == "High" else GRAY
    run = p.add_run(priority)
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = priority_color
    run.bold = (priority in ("Critical", "High"))

# ── NOT FIXABLE ON SQUARESPACE ──
sub_heading(doc, "\u26D4  Platform Limitations (Cannot Be Fixed on Squarespace)")

add_text(doc, "The following issues are caused by how Squarespace builds and serves pages. "
         "No amount of optimization within the platform can resolve them. These are the reasons "
         "the site\u2019s SEO score has a hard ceiling on Squarespace.",
         size=9, color=GRAY, italic=True, space_before=2, space_after=10)

unfixable_items = [
    ("Low text-to-HTML ratio (47 pages)",
     "Squarespace injects 20+ CSS and JavaScript resources on every page \u2014 analytics, "
     "font loaders, animation libraries, commerce scripts, and platform tracking \u2014 whether "
     "you use them or not. This bloats the code-to-content ratio well below the 25% threshold "
     "search engines prefer. On a custom site, you only load what you need. On Squarespace, "
     "you carry the weight of the entire platform on every page.",
     "This is the single biggest reason Squarespace sites score lower on technical SEO. "
     "Google\u2019s crawlers have a limited budget for each site, and when most of that budget "
     "is spent parsing unnecessary code instead of reading your content, your pages rank lower."),
    ("Missing or duplicate H1 headings (16 missing, 9 multiple)",
     "Squarespace auto-generates H1 tags from page titles and sometimes from template elements "
     "you can\u2019t control. Some templates produce no H1 at all; others produce multiple H1s. "
     "You can manually set text to H1 in the editor, but the template may still inject its own, "
     "creating conflicts. There\u2019s no way to reliably guarantee one H1 per page without "
     "custom code injection.",
     "H1 tags tell Google what a page is about. Multiple H1s dilute that signal. Missing H1s "
     "mean Google has to guess. Both hurt rankings for your target keywords."),
    ("Page speed / TTFB (Time to First Byte)",
     "Squarespace serves pages from shared infrastructure. You cannot control server response "
     "times, caching headers, CDN configuration, or asset compression. Page load speeds are "
     "consistently slower than custom-hosted sites. Squarespace\u2019s average TTFB is "
     "800\u20131500ms; a well-configured custom site achieves 50\u2013200ms.",
     "Google uses page speed as a direct ranking factor (Core Web Vitals). Slower sites rank "
     "lower, and visitors leave before the page finishes loading \u2014 increasing bounce rate, "
     "which further suppresses rankings."),
    ("No custom robots.txt or advanced sitemap control",
     "Squarespace auto-generates robots.txt and sitemap.xml. You cannot customize which pages "
     "are included/excluded, add sitemap priorities, or block specific crawlers. If Squarespace "
     "includes pages you don\u2019t want indexed, your only option is the \u201cHide from search\u201d "
     "toggle per page, which uses a noindex meta tag rather than proper robots.txt directives.",
     "Without fine-grained crawl control, search engine bots waste time on pages that don\u2019t "
     "matter \u2014 login pages, empty tag pages, utility URLs \u2014 instead of spending their "
     "limited crawl budget on your money pages."),
    ("No custom structured data (JSON-LD / Schema)",
     "Squarespace injects its own limited schema markup (basic Organization, Website). You cannot "
     "add custom LocalBusiness schema, FAQPage schema, Service schema, or any structured data "
     "that Google uses for rich results. Code injection blocks can add JSON-LD to headers, but "
     "it\u2019s fragile and limited to site-wide injection, not per-page.",
     "Without proper schema, your business won\u2019t appear in Google\u2019s knowledge panel, "
     "rich snippets, or AI-powered search results. Your competitors with custom sites get "
     "enhanced listings while you show up as plain blue links."),
    ("No llms.txt or AI search optimization",
     "Squarespace does not support creating custom files at the root domain like llms.txt. "
     "This file tells AI search engines (ChatGPT, Perplexity, Google AI Overviews) how to "
     "understand and cite your site. Without it, AI assistants may ignore or misrepresent "
     "your business when answering patient queries.",
     "AI-powered search is growing rapidly. Businesses without llms.txt and speakable schema "
     "are invisible to the next generation of search."),
    ("Pages blocked from crawling (platform-generated)",
     "Squarespace blocks certain system pages from crawling by default. You cannot unblock "
     "these or control which platform-generated URLs are accessible to search engines.",
     "Not usually harmful, but removes control over what Google sees and indexes."),
]

for name, desc, impact in unfixable_items:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(6.5)
    remove_table_borders(table)
    set_cell_borders(table.cell(0, 0), bottom={'sz': '2', 'color': 'DDDDDD'})

    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"\u26D4  {name}")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    run.bold = True

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(desc)
    run2.font.name = "Arial"
    run2.font.size = Pt(8.5)
    run2.font.color.rgb = DARK_GRAY

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(8)
    run3 = p3.add_run(f"Business impact: {impact}")
    run3.font.name = "Arial"
    run3.font.size = Pt(8.5)
    run3.font.color.rgb = RGBColor(0x88, 0x33, 0x33)
    run3.bold = True

# ── BOTTOM LINE ──
add_spacer(doc, 8)

sub_heading(doc, "The Bottom Line")

add_text(doc, "By fixing everything in the \u2705 Fixable list, we can push the site\u2019s "
         "SEO health score from ___% to approximately ___%  \u2014 a meaningful improvement that "
         "will directly increase rankings and phone calls. But the \u26D4 Platform Limitations "
         "create a ceiling that no amount of Squarespace optimization can break through.",
         size=10, color=DARK_GRAY, space_before=4, space_after=8)

add_text(doc, "A custom-built site on a modern framework (Next.js, Astro, or similar) with proper "
         "hosting would eliminate every item in the unfixable list and put the site on track for "
         "a 95%+ SEO health score. That means faster load times, richer Google listings, AI search "
         "visibility, and no wasted crawl budget \u2014 all of which translate directly to more "
         "patients finding Form Chiropractic when they search.",
         size=10, color=DARK_GRAY, space_before=0, space_after=8)

add_text(doc, "We recommend maximizing what\u2019s possible on Squarespace now (the fixable items "
         "above are all included in your current $200/mo SEO package) while planning a migration "
         "to a custom site when budget allows. The SEO work we do now \u2014 content, backlinks, "
         "GBP optimization \u2014 transfers to a new site. Nothing is wasted.",
         size=10, color=DARK_GRAY, space_before=0, space_after=12)


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE & RE-INJECT LOGO
# ═══════════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT_DOCX)

# Re-inject official logo
buf = io.BytesIO()
with ZipFile(OUTPUT_DOCX, 'r') as zin, ZipFile(buf, 'w') as zout:
    for item in zin.namelist():
        if item.startswith('word/media/'):
            with open(LOGO_PATH, 'rb') as f:
                zout.writestr(item, f.read())
        else:
            zout.writestr(item, zin.read(item))
with open(OUTPUT_DOCX, 'wb') as f:
    f.write(buf.getvalue())

print(f"✅ Template saved: {OUTPUT_DOCX}")
